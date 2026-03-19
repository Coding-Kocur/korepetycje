import os
from docx import Document
from docx.shared import Pt, RGBColor

def create_client_materials(topic, path, warmup, videos, grammar_title, grammar_rules, vocab, reading_info, questions):
    doc = Document()
    doc.add_heading(f'ESL Lesson: {topic}', 0)
    
    # Warm-up
    doc.add_heading('Warm-up', level=1)
    for i, q in enumerate(warmup, 1):
        doc.add_paragraph(f'{i}. {q}')
        
    # Video Analysis
    doc.add_heading('Video Analysis', level=1)
    for v in videos:
        doc.add_paragraph(f"{v['name']} Report: {v['url']}")
        
    for v in videos:
        doc.add_heading(f"{v['name']} Questions", level=2)
        for i, q in enumerate(v['questions'], 1):
            doc.add_paragraph(f'{i}. {q}')
            
    # Grammar
    doc.add_heading('Grammar: ' + grammar_title, level=1)
    doc.add_paragraph(grammar_rules)
    doc.add_heading('Grammar Exercises', level=2)
    for i, q in enumerate(questions['grammar'], 1):
        doc.add_paragraph(f'{i}. {q}')
        
    # Vocabulary
    doc.add_heading('Vocabulary', level=1)
    for item in vocab:
        doc.add_paragraph(item)
        
    # Reading
    doc.add_heading('Reading Section', level=1)
    doc.add_paragraph(reading_info['text'])
    p = doc.add_paragraph()
    r = p.add_run(f"Source: {reading_info['source']} | Author: {reading_info['author']} | Date: {reading_info['date']}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_heading('Reading Questions', level=2)
    for i, q in enumerate(questions['reading'], 1):
        doc.add_paragraph(f'{i}. {q}')
        
    doc.save(os.path.join(path, 'Client_Materials.docx'))


def create_teacher_guide(topic, path, warmup, videos, grammar_title, grammar_rules, vocab, reading_info, questions, answers):
    doc = Document()
    doc.add_heading(f'TEACHER GUIDE: {topic}', 0)
    
    # Video Timestamps
    doc.add_heading('Video Timestamps', level=1)
    for v in videos:
        doc.add_paragraph(f"{v['name']}: Focus on {v['focus']}")
        
    # Warm-up
    doc.add_heading('Warm-up', level=1)
    for i, q in enumerate(warmup, 1):
        doc.add_paragraph(f'{i}. {q}')

    # Video Analysis Answers
    doc.add_heading('Video Analysis', level=1)
    for v in videos:
        doc.add_paragraph(f"{v['name']} Report: {v['url']}")
    
    for v in videos:
        doc.add_heading(f"{v['name']} Questions", level=2)
        for i, (q, ans) in enumerate(zip(v['questions'], answers['video'][v['name']]), 1):
            doc.add_paragraph(f'{i}. {q}')
            p = doc.add_paragraph(f'Expected [{ans["ts"]}]: {ans["text"]}')
            p.runs[0].font.color.rgb = RGBColor(0, 100, 0)
            
    # Grammar Answers
    doc.add_heading('Grammar: ' + grammar_title, level=1)
    doc.add_heading('Grammar Exercises (Answers)', level=2)
    for i, (q, a) in enumerate(zip(questions['grammar'], answers['grammar']), 1):
        doc.add_paragraph(f'{i}. {q}')
        p = doc.add_paragraph(f'Answer: {a}')
        p.runs[0].font.color.rgb = RGBColor(0, 100, 0)
        
    # Reading with underlines
    doc.add_heading('Reading Section (Answers in Text)', level=1)
    text = reading_info['text']
    p = doc.add_paragraph()
    
    # Simple logic to underline target sentences
    underlines = answers['reading_underlines']
    
    # Split the text by sentences roughly
    sentences = text.split('. ')
    for i, sentence in enumerate(sentences):
        if i > 0:
            p.add_run('. ')
        
        is_underlined = any(u in sentence for u in underlines)
        run = p.add_run(sentence)
        if is_underlined:
            run.underline = True
            run.font.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            
    # Reading Questions
    doc.add_heading('Reading Questions (Answers)', level=2)
    for i, (q, a) in enumerate(zip(questions['reading'], answers['reading']), 1):
        doc.add_paragraph(f'{i}. {q}')
        p = doc.add_paragraph(f'Answer: {a}')
        p.runs[0].font.color.rgb = RGBColor(0, 100, 0)

    doc.save(os.path.join(path, 'Teacher_Guide.docx'))

# DATA FOR EV
ev_path = 'Electric Cars A2'
os.makedirs(ev_path, exist_ok=True)
ev_warmup = [
    "Do you drive a car? What kind of fuel does it use?",
    "Do you think electric cars are the future of transport?",
    "What are some problems with electric cars today?",
    "Would you like to buy an electric car in the next 5 years?"
]
ev_videos = [
    {
        "name": "Network A", 
        "url": "https://www.youtube.com/watch?v=kzsyJROQOas", 
        "focus": "1:00 - 2:30",
        "questions": [
            "What replaces the gas tank in an electric car?",
            "How do people usually charge electric cars at home?",
            "What happens when you press the pedal in an electric car?"
        ]
    },
    {
        "name": "Network B", 
        "url": "https://www.youtube.com/watch?v=cT2-AJj0yfo", 
        "focus": "0:30 - 2:00",
        "questions": [
            "What is the main advantage of an electric motor?",
            "Why is the battery pack usually placed at the bottom of the car?",
            "Does an electric car have an exhaust pipe?"
        ]
    }
]
ev_grammar_rules = "FIRST CONDITIONAL: We use the First Conditional to talk about a possible future action.\nRule: If + Present Simple, will + verb.\nExamples:\n- If you buy an EV, you will try to save money on gasoline.\n- If it rains, I will stay home."
ev_questions = {
    "grammar": [
        "If my car (break) down, I (call) a mechanic.",
        "If you (drive) an electric car, you (not produce) exhaust gases.",
        "We (spend) less money if we (charge) the car at home."
    ],
    "reading": [
        "How many electric cars were sold globally in 2024?",
        "Which country had the most electric car sales?",
        "What problem did some companies see in early 2026?"
    ]
}
ev_vocab = [
    "Battery - A device that stores electricity.",
    "Engine - A machine that makes a car move.",
    "Charge - To put electricity into a battery.",
    "Distance - How far something is.",
    "Exhaust - The dirty gas that comes out of a car.",
    "Sales - The number of things that are sold.",
    "Growth - The process of becoming larger or more.",
    "Global - Affecting the whole world.",
    "Subsidies - Money from the government to help pay for something.",
    "Consumer - A person who buys things."
]
ev_reading = {
    "text": "Global electric vehicle sales reached a new milestone in 2024. More than 17 million units were sold, which was a 25 percent increase from the previous year. China was the biggest market, selling over 10 million electric cars. Many people in Europe and North America also bought electric cars. However, in early 2026, some companies saw a slowdown in sales. Experts say this is because governments reduced financial subsidies and because some drivers are worried about battery charging options on long journeys. If governments build more public chargers, more consumers will likely choose electric vehicles.",
    "source": "Adapted from Reuters Reports",
    "author": "News Team",
    "date": "March 2024"
}
ev_answers = {
    "video": {
        "Network A": [
            {"ts": "1:15", "text": "A large battery pack replaces the gas tank."},
            {"ts": "1:40", "text": "They plug it into the wall overnight."},
            {"ts": "2:10", "text": "Electricity goes directly to the motor and the car moves instantly."}
        ],
        "Network B": [
            {"ts": "0:45", "text": "It is very quiet and smooth."},
            {"ts": "1:05", "text": "To keep the car balanced and safe on the road."},
            {"ts": "1:30", "text": "No, it does not have an exhaust pipe because there are no emissions."}
        ]
    },
    "grammar": [
        "breaks / will call",
        "drive / will not produce",
        "will spend / charge"
    ],
    "reading_underlines": [
        "More than 17 million units were sold",
        "China was the biggest market",
        "in early 2026, some companies saw a slowdown in sales"
    ],
    "reading": [
        "More than 17 million units were sold.",
        "China was the biggest market.",
        "They saw a slowdown in sales."
    ]
}

create_client_materials('Electric Cars and Global Shift (A2)', ev_path, ev_warmup, ev_videos, 'First Conditional', ev_grammar_rules, ev_vocab, ev_reading, ev_questions)
create_teacher_guide('Electric Cars and Global Shift (A2)', ev_path, ev_warmup, ev_videos, 'First Conditional', ev_grammar_rules, ev_vocab, ev_reading, ev_questions, ev_answers)

# DATA FOR SPACE
space_path = 'Space Exploration A2'
os.makedirs(space_path, exist_ok=True)
sp_warmup = [
    "Have you ever looked at the Moon through a telescope?",
    "Would you like to travel to space as a tourist?",
    "Why do you think countries spend money on exploring space?",
    "Do you think humans will live on another planet in the future?"
]
sp_videos = [
    {
        "name": "Network A", 
        "url": "https://www.youtube.com/watch?v=_T8cn2J13-4", 
        "focus": "1:00 - 2:30",
        "questions": [
            "What is the name of the new NASA rocket?",
            "What is the ultimate goal after returning to the Moon?",
            "When will astronauts step on the Moon again?"
        ]
    },
    {
        "name": "Network B", 
        "url": "https://www.youtube.com/watch?v=-oaJmgm8hs4", 
        "focus": "1:30 - 3:00",
        "questions": [
            "What year did Apollo 11 land on the Moon?",
            "What was the name of the lunar module?",
            "How long did the trip from Earth to the Moon take?"
        ]
    }
]
sp_grammar_rules = "FIRST CONDITIONAL: We use the First Conditional to talk about a possible future action.\nRule: If + Present Simple, will + verb.\nExamples:\n- If SpaceX builds cheap rockets, we will travel to space more often.\n- If astronauts find water on the Moon, they will stay longer."
sp_questions = {
    "grammar": [
        "If NASA (find) water on the Moon, they (build) a base there.",
        "If you (not wear) a spacesuit in space, you (not survive).",
        "The mission (be) successful if the rocket (launch) safely."
    ],
    "reading": [
        "When is SpaceX planning to send cargo to the Moon?",
        "What is the name of the SpaceX rocket that will carry astronauts?",
        "Why is the Moon better than Mars for building a first city?"
    ]
}
sp_vocab = [
    "Rocket - A vehicle used for traveling into space.",
    "Astronaut - A person who goes into space.",
    "Orbit - The path of an object moving around a planet.",
    "Lunar - Related to the Moon.",
    "Planet - A large round object in space that circles a star.",
    "Gravity - The invisible force that pulls things to the ground.",
    "Base - A place where people live and work.",
    "Cargo - Goods carried by a ship, plane, or spacecraft.",
    "Explore - To travel around a new place to learn about it.",
    "Distance - How far away something is."
]
sp_reading = {
    "text": "SpaceX is making big plans for the Moon. The company wants to help send astronauts back to the lunar surface. They are using a new massive vehicle called Starship. If everything goes well, SpaceX aims to start sending cargo flights to the Moon by 2028. The CEO, Elon Musk, thinks we can build a self-growing city on the Moon much sooner than on Mars. The Moon is much closer to Earth, which allows for faster trips and quicker progress. This new step in space exploration will help scientists learn how to live in space before attempting the long trip to Mars.",
    "source": "Adapted from Reuters and Space.com",
    "author": "Space News Desk",
    "date": "March 2024"
}
sp_answers = {
    "video": {
        "Network A": [
            {"ts": "1:20", "text": "The new rocket is called SLS (Space Launch System)."},
            {"ts": "2:05", "text": "The final goal is to send humans to Mars."},
            {"ts": "2:25", "text": "They plan to step on the Moon again later this decade."}
        ],
        "Network B": [
            {"ts": "1:40", "text": "Apollo 11 landed on the Moon in 1969."},
            {"ts": "2:15", "text": "It was called the Eagle."},
            {"ts": "2:50", "text": "The trip took about three days."}
        ]
    },
    "grammar": [
        "finds / will build",
        "do not wear / will not survive",
        "will be / launches"
    ],
    "reading_underlines": [
        "SpaceX aims to start sending cargo flights to the Moon by 2028",
        "They are using a new massive vehicle called Starship",
        "The Moon is much closer to Earth, which allows for faster trips"
    ],
    "reading": [
        "They want to send cargo flights by 2028.",
        "It is called Starship.",
        "Because it is closer, allowing for faster trips."
    ]
}

create_client_materials('Return to the Moon (A2)', space_path, sp_warmup, sp_videos, 'First Conditional', sp_grammar_rules, sp_vocab, sp_reading, sp_questions)
create_teacher_guide('Return to the Moon (A2)', space_path, sp_warmup, sp_videos, 'First Conditional', sp_grammar_rules, sp_vocab, sp_reading, sp_questions, sp_answers)

print("Created all files successfully.")
